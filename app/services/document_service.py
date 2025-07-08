import os
import json
import re
from datetime import datetime
from flask import current_app
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import black
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image
import io
import base64


class DocumentService:
    """文档服务 - 处理文档预览和导出"""
    
    # 支持预览的文件类型
    PREVIEW_SUPPORTED_TYPES = {
        'pdf', 'txt', 'md', 'json', 'xml', 'doc', 'docx', 
        'png', 'jpg', 'jpeg', 'gif', 'bmp'
    }
    
    @staticmethod
    def can_preview(file_path):
        """检查文件是否支持预览"""
        try:
            if not os.path.exists(file_path):
                return False
            
            file_extension = os.path.splitext(file_path)[1].lower().lstrip('.')
            return file_extension in DocumentService.PREVIEW_SUPPORTED_TYPES
        except:
            return False
    
    @staticmethod
    def get_file_preview(file_path, max_length=5000):
        """获取文件预览内容"""
        try:
            if not os.path.exists(file_path):
                return {'error': '文件不存在', 'content': None}
            
            file_extension = os.path.splitext(file_path)[1].lower().lstrip('.')
            
            if file_extension == 'pdf':
                return DocumentService._preview_pdf(file_path, max_length)
            elif file_extension in ['txt', 'md']:
                return DocumentService._preview_text(file_path, max_length)
            elif file_extension == 'json':
                return DocumentService._preview_json(file_path, max_length)
            elif file_extension == 'xml':
                return DocumentService._preview_xml(file_path, max_length)
            elif file_extension in ['doc', 'docx']:
                return DocumentService._preview_docx(file_path, max_length)
            elif file_extension in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
                return DocumentService._preview_image(file_path)
            else:
                return {'error': '不支持预览此文件类型', 'content': None}
                
        except Exception as e:
            current_app.logger.error(f"文件预览失败: {file_path} - 错误: {str(e)}", exc_info=True)
            return {'error': f'预览失败: {str(e)}', 'content': None}
    
    @staticmethod
    def _preview_pdf(file_path, max_length):
        """预览PDF文件"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                content = []
                total_length = 0
                
                for page_num, page in enumerate(pdf_reader.pages):
                    if total_length >= max_length:
                        break
                    
                    page_text = page.extract_text()
                    if page_text:
                        remaining_length = max_length - total_length
                        if len(page_text) > remaining_length:
                            page_text = page_text[:remaining_length] + "..."
                        
                        content.append(f"--- 第 {page_num + 1} 页 ---\n{page_text}")
                        total_length += len(page_text)
                
                return {
                    'content': '\n\n'.join(content),
                    'type': 'text',
                    'pages': len(pdf_reader.pages),
                    'truncated': total_length >= max_length
                }
                
        except Exception as e:
            return {'error': f'PDF解析失败: {str(e)}', 'content': None}
    
    @staticmethod
    def _preview_text(file_path, max_length):
        """预览文本文件"""
        try:
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read(max_length)
                        truncated = len(content) == max_length
                        
                        return {
                            'content': content,
                            'type': 'text',
                            'encoding': encoding,
                            'truncated': truncated
                        }
                except UnicodeDecodeError:
                    continue
            
            return {'error': '无法解码文件，不支持的编码格式', 'content': None}
            
        except Exception as e:
            return {'error': f'文本读取失败: {str(e)}', 'content': None}
    
    @staticmethod
    def _preview_json(file_path, max_length):
        """预览JSON文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                
            # 格式化JSON
            formatted_json = json.dumps(data, ensure_ascii=False, indent=2)
            
            if len(formatted_json) > max_length:
                formatted_json = formatted_json[:max_length] + "\n... (内容已截断)"
                truncated = True
            else:
                truncated = False
            
            return {
                'content': formatted_json,
                'type': 'json',
                'truncated': truncated
            }
            
        except Exception as e:
            return {'error': f'JSON解析失败: {str(e)}', 'content': None}
    
    @staticmethod
    def _preview_xml(file_path, max_length):
        """预览XML文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read(max_length)
                truncated = len(content) == max_length
                
                return {
                    'content': content,
                    'type': 'xml',
                    'truncated': truncated
                }
                
        except Exception as e:
            return {'error': f'XML读取失败: {str(e)}', 'content': None}
    
    @staticmethod
    def _preview_docx(file_path, max_length):
        """预览DOCX文件"""
        try:
            doc = DocxDocument(file_path)
            
            content = []
            total_length = 0
            
            for para in doc.paragraphs:
                if total_length >= max_length:
                    break
                
                text = para.text.strip()
                if text:
                    remaining_length = max_length - total_length
                    if len(text) > remaining_length:
                        text = text[:remaining_length] + "..."
                    
                    content.append(text)
                    total_length += len(text)
            
            return {
                'content': '\n\n'.join(content),
                'type': 'text',
                'paragraphs': len(doc.paragraphs),
                'truncated': total_length >= max_length
            }
            
        except Exception as e:
            return {'error': f'DOCX解析失败: {str(e)}', 'content': None}
    
    @staticmethod
    def _preview_image(file_path):
        """预览图片文件"""
        try:
            with Image.open(file_path) as img:
                # 获取图片信息
                width, height = img.size
                format_name = img.format
                mode = img.mode
                
                # 如果图片太大，创建缩略图
                max_size = (800, 600)
                if width > max_size[0] or height > max_size[1]:
                    img.thumbnail(max_size, Image.Resampling.LANCZOS)
                
                # 转换为base64
                buffer = io.BytesIO()
                img.save(buffer, format='PNG')
                img_data = buffer.getvalue()
                img_base64 = base64.b64encode(img_data).decode('utf-8')
                
                return {
                    'content': f"data:image/png;base64,{img_base64}",
                    'type': 'image',
                    'width': width,
                    'height': height,
                    'format': format_name,
                    'mode': mode
                }
                
        except Exception as e:
            return {'error': f'图片处理失败: {str(e)}', 'content': None}
    

    
    @staticmethod
    def export_task_result_to_pdf(task_result, output_path=None):
        """导出任务结果为PDF - 支持Markdown预览格式转换"""
        try:
            if not output_path:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = f"task_result_{task_result.task_id}_{timestamp}.pdf"
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # 新的实现：优先使用Markdown转HTML转PDF方式
            success, pdf_path, error = DocumentService._export_markdown_to_pdf(
                task_result, output_path
            )
            
            if success:
                try:
                    current_app.logger.info(f"PDF导出成功（Markdown预览格式）: {output_path}")
                except RuntimeError:
                    print(f"PDF导出成功（Markdown预览格式）: {output_path}")
                return True, pdf_path, None
            else:
                # 如果新方法失败，回退到原有方法
                try:
                    current_app.logger.warning(f"Markdown转PDF失败，使用原有方法: {error}")
                except RuntimeError:
                    print(f"Markdown转PDF失败，使用原有方法: {error}")
                
                return DocumentService._export_text_to_pdf_legacy(task_result, output_path)
                
        except Exception as e:
            error_msg = f"PDF导出失败: {str(e)}"
            try:
                current_app.logger.error(error_msg, exc_info=True)
            except RuntimeError:
                print(error_msg)
                import traceback
                traceback.print_exc()
            return False, None, error_msg
    
    @staticmethod
    def _export_markdown_to_pdf(task_result, output_path):
        """新方法：将Markdown内容转换为HTML，然后转换为PDF"""
        try:
            # 获取任务结果内容
            content = task_result.answer or "暂无处理结果"
            content = str(content).strip()
            
            # 清理和转换Markdown内容
            cleaned_content = DocumentService._clean_markdown_content(content)
            html_content = DocumentService._convert_markdown_to_html(cleaned_content)
            
            # 创建专门用于PDF的HTML页面（简化样式，适合打印）
            full_html = DocumentService._create_pdf_html_page(
                html_content, 
                f"任务结果 - {task_result.task_id}"
            )
            
            # 使用weasyprint将HTML转换为PDF
            try:
                from weasyprint import HTML, CSS
                from weasyprint.text.fonts import FontConfiguration
                
                # 创建字体配置
                font_config = FontConfiguration()
                
                # 生成PDF
                html_doc = HTML(string=full_html)
                html_doc.write_pdf(output_path, font_config=font_config)
                
                return True, output_path, None
                
            except ImportError as import_error:
                return False, None, f"weasyprint库未安装: {str(import_error)}"
            except OSError as os_error:
                if "gobject" in str(os_error) or "GTK" in str(os_error):
                    return False, None, f"weasyprint需要GTK库支持（Windows系统常见问题）: {str(os_error)}"
                else:
                    return False, None, f"weasyprint系统库依赖问题: {str(os_error)}"
            except Exception as pdf_error:
                return False, None, f"HTML转PDF失败: {str(pdf_error)}"
                
        except Exception as e:
            return False, None, f"Markdown预处理失败: {str(e)}"
    
    @staticmethod
    def _clean_markdown_content(content):
        """清理和格式化Markdown内容，处理数据库中存储的格式"""
        if not content:
            return "# 任务结果\n\n暂无处理结果"
        
        # 确保是字符串
        content = str(content)
        original_content = content
        
        # 移除外层的代码块标记（处理数据库中存储的格式）
        content_stripped = content.strip()
        
        try:
            current_app.logger.info(f"原始内容前100字符: {content_stripped[:100]}...")
            current_app.logger.info(f"内容开头检查: startswith('```markdown'): {content_stripped.startswith('```markdown')}")
            current_app.logger.info(f"内容结尾检查: endswith('```'): {content_stripped.endswith('```')}")
        except RuntimeError:
            print(f"原始内容前100字符: {content_stripped[:100]}...")
            print(f"内容开头检查: startswith('```markdown'): {content_stripped.startswith('```markdown')}")
            print(f"内容结尾检查: endswith('```'): {content_stripped.endswith('```')}")
        
        # 处理各种代码块包装情况（改进版本，更robust）
        cleaned = False
        
        # 检查````markdown ... ````格式
        if content_stripped.startswith('````markdown') and content_stripped.endswith('````'):
            content = content_stripped[12:-4].strip()
            cleaned = True
            try:
                current_app.logger.info("✅ 检测到````markdown代码块格式，已自动提取内容")
            except RuntimeError:
                print("✅ 检测到````markdown代码块格式，已自动提取内容")
                
        # 检查```markdown ... ```格式
        elif content_stripped.startswith('```markdown') and content_stripped.endswith('```'):
            content = content_stripped[11:-3].strip()
            cleaned = True
            try:
                current_app.logger.info("✅ 检测到```markdown代码块格式，已自动提取内容")
            except RuntimeError:
                print("✅ 检测到```markdown代码块格式，已自动提取内容")
                
        # 检查````...````格式
        elif content_stripped.startswith('````') and content_stripped.endswith('````'):
            content = content_stripped[4:-4].strip()
            cleaned = True
            try:
                current_app.logger.info("✅ 检测到````代码块格式，已自动提取内容")
            except RuntimeError:
                print("✅ 检测到````代码块格式，已自动提取内容")
                
        # 检查```...```格式（逐行处理，更准确）
        elif content_stripped.startswith('```') and content_stripped.endswith('```'):
            lines = content_stripped.split('\n')
            if len(lines) >= 2:
                first_line = lines[0].strip()
                last_line = lines[-1].strip()
                
                # 检查第一行是否是```或```语言名
                if first_line.startswith('```') and last_line == '```':
                    content = '\n'.join(lines[1:-1]).strip()
                    cleaned = True
                    try:
                        current_app.logger.info(f"✅ 检测到```代码块格式（首行：{first_line}），已自动提取内容")
                    except RuntimeError:
                        print(f"✅ 检测到```代码块格式（首行：{first_line}），已自动提取内容")
        
        # 如果没有清理，记录日志
        if not cleaned:
            try:
                current_app.logger.info("ℹ️ 内容格式不匹配任何代码块模式，保持原样")
            except RuntimeError:
                print("ℹ️ 内容格式不匹配任何代码块模式，保持原样")
        
        # 基本清理
        content = content.strip()
        
        # 移除元数据信息（文档类型、转换时间、源格式等）
        content = DocumentService._remove_metadata_info(content)
        
        # 如果清理后内容为空，提供默认内容
        if not content:
            return "# 任务结果\n\n暂无处理结果"
        
        # 如果内容不是以标题开始，添加一个主标题
        if not content.startswith('#'):
            content = "# 任务结果\n\n" + content
        
        # 确保段落之间有适当的间距
        lines = content.split('\n')
        cleaned_lines = []
        prev_line_empty = False
        
        for line in lines:
            line = line.rstrip()
            is_empty = len(line) == 0
            
            # 避免连续的空行
            if is_empty and prev_line_empty:
                continue
                
            cleaned_lines.append(line)
            prev_line_empty = is_empty
        
        cleaned_content = '\n'.join(cleaned_lines)
        
        # 验证最终结果
        still_has_codeblock = cleaned_content.strip().startswith('```') or cleaned_content.strip().endswith('```')
        
        try:
            current_app.logger.info(f"清理后内容前100字符: {cleaned_content[:100]}...")
            current_app.logger.info(f"清理结果验证: 是否还有代码块标记: {still_has_codeblock}")
            if still_has_codeblock:
                current_app.logger.warning("⚠️ 警告：清理后内容仍包含代码块标记，可能需要手动检查")
        except RuntimeError:
            print(f"清理后内容前100字符: {cleaned_content[:100]}...")
            print(f"清理结果验证: 是否还有代码块标记: {still_has_codeblock}")
            if still_has_codeblock:
                print("⚠️ 警告：清理后内容仍包含代码块标记，可能需要手动检查")
        
        return cleaned_content
    
    @staticmethod
    def _remove_metadata_info(content):
        """移除内容中的元数据信息（文档类型、转换时间、源格式等）"""
        if not content:
            return content
        
        lines = content.split('\n')
        cleaned_lines = []
        skip_metadata_section = False
        
        # 检测和移除元数据模式
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # 检测元数据开始标记
            metadata_patterns = [
                '> **文档类型**：',
                '> **转换时间**：', 
                '> **源格式**：',
                '**文档类型**：',
                '**转换时间**：',
                '**源格式**：',
                '文档类型：',
                '转换时间：',
                '源格式：'
            ]
            
            # 检查是否是元数据行
            is_metadata_line = any(pattern in line_stripped for pattern in metadata_patterns)
            
            # 检查是否是元数据分隔线
            is_separator_line = line_stripped in ['---', '***', '___'] or line_stripped.startswith('---')
            
            # 如果遇到元数据行，开始跳过模式
            if is_metadata_line:
                skip_metadata_section = True
                try:
                    current_app.logger.info(f"检测到元数据行，开始跳过: {line_stripped[:50]}...")
                except RuntimeError:
                    print(f"检测到元数据行，开始跳过: {line_stripped[:50]}...")
                continue
            
            # 如果在跳过模式中遇到分隔线，结束跳过模式
            if skip_metadata_section and is_separator_line:
                skip_metadata_section = False
                try:
                    current_app.logger.info("检测到元数据分隔线，结束跳过模式")
                except RuntimeError:
                    print("检测到元数据分隔线，结束跳过模式")
                continue
            
            # 如果在跳过模式中，继续跳过
            if skip_metadata_section:
                continue
            
            # 保留非元数据行
            cleaned_lines.append(line)
        
        # 重新组合内容
        cleaned_content = '\n'.join(cleaned_lines)
        
        # 移除开头的多余空行
        cleaned_content = cleaned_content.lstrip('\n')
        
        try:
            current_app.logger.info(f"元数据清理完成，清理前行数: {len(lines)}, 清理后行数: {len(cleaned_lines)}")
        except RuntimeError:
            print(f"元数据清理完成，清理前行数: {len(lines)}, 清理后行数: {len(cleaned_lines)}")
        
        return cleaned_content
    
    @staticmethod
    def _convert_html_to_reportlab_elements(html_content, base_style, chinese_available, chinese_font_name):
        """将HTML内容转换为ReportLab元素列表，实现预览格式渲染"""
        from bs4 import BeautifulSoup
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        
        elements = []
        
        try:
            # 解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 创建样式
            if chinese_available:
                # 基本样式
                normal_style = ParagraphStyle(
                    'ChineseNormal',
                    parent=base_style,
                    fontName=chinese_font_name,
                    fontSize=12,
                    spaceAfter=12,
                    leftIndent=0
                )
                
                # 标题样式
                h1_style = ParagraphStyle(
                    'ChineseH1',
                    parent=base_style,
                    fontName=chinese_font_name,
                    fontSize=18,
                    spaceAfter=16,
                    spaceBefore=16,
                    textColor=colors.black
                )
                
                h2_style = ParagraphStyle(
                    'ChineseH2',
                    parent=base_style,
                    fontName=chinese_font_name,
                    fontSize=16,
                    spaceAfter=14,
                    spaceBefore=14,
                    textColor=colors.black
                )
                
                h3_style = ParagraphStyle(
                    'ChineseH3',
                    parent=base_style,
                    fontName=chinese_font_name,
                    fontSize=14,
                    spaceAfter=12,
                    spaceBefore=12,
                    textColor=colors.black
                )
                
                # 引用样式
                blockquote_style = ParagraphStyle(
                    'ChineseBlockquote',
                    parent=base_style,
                    fontName=chinese_font_name,
                    fontSize=11,
                    spaceAfter=12,
                    leftIndent=20,
                    textColor=colors.grey
                )
                
                # 代码样式
                code_style = ParagraphStyle(
                    'ChineseCode',
                    parent=base_style,
                    fontName='Courier',  # 代码使用等宽字体
                    fontSize=10,
                    spaceAfter=8,
                    leftIndent=10,
                    backColor=colors.lightgrey
                )
            else:
                # 非中文字体的备选样式
                normal_style = base_style  
                h1_style = base_style
                h2_style = base_style
                h3_style = base_style
                blockquote_style = base_style
                code_style = base_style
            
            # 处理HTML元素
            for element in soup.body.children if soup.body else soup.children:
                if hasattr(element, 'name'):
                    if element.name == 'h1':
                        text = element.get_text().strip()
                        if text:
                            elements.append(Paragraph(text, h1_style))
                            elements.append(Spacer(1, 6))
                            
                    elif element.name == 'h2':
                        text = element.get_text().strip()
                        if text:
                            elements.append(Paragraph(text, h2_style))
                            elements.append(Spacer(1, 4))
                            
                    elif element.name == 'h3':
                        text = element.get_text().strip()
                        if text:
                            elements.append(Paragraph(text, h3_style))
                            elements.append(Spacer(1, 4))
                            
                    elif element.name == 'p':
                        text = element.get_text().strip()
                        if text:
                            # 处理段落中的内联元素
                            formatted_text = DocumentService._process_inline_html_elements(element)
                            elements.append(Paragraph(formatted_text, normal_style))
                            elements.append(Spacer(1, 6))
                            
                    elif element.name == 'blockquote':
                        text = element.get_text().strip()
                        if text:
                            # 引用块
                            elements.append(Paragraph(f"> {text}", blockquote_style))
                            elements.append(Spacer(1, 6))
                            
                    elif element.name == 'ul':
                        # 无序列表
                        for li in element.find_all('li'):
                            text = li.get_text().strip()
                            if text:
                                elements.append(Paragraph(f"• {text}", normal_style))
                        elements.append(Spacer(1, 6))
                        
                    elif element.name == 'ol':
                        # 有序列表
                        for i, li in enumerate(element.find_all('li'), 1):
                            text = li.get_text().strip()
                            if text:
                                elements.append(Paragraph(f"{i}. {text}", normal_style))
                        elements.append(Spacer(1, 6))
                        
                    elif element.name == 'table':
                        # 表格处理
                        table_data = []
                        for row in element.find_all('tr'):
                            row_data = []
                            for cell in row.find_all(['td', 'th']):
                                cell_text = cell.get_text().strip()
                                row_data.append(cell_text)
                            if row_data:
                                table_data.append(row_data)
                        
                        if table_data:
                            # 创建表格
                            table = Table(table_data)
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                                ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke),
                                ('ALIGN',(0,0),(-1,-1),'CENTER'),
                                ('FONTNAME', (0,0), (-1,0), chinese_font_name if chinese_available else 'Helvetica-Bold'),
                                ('FONTSIZE', (0,0), (-1,0), 12),
                                ('BOTTOMPADDING', (0,0), (-1,0), 12),
                                ('BACKGROUND',(0,1),(-1,-1),colors.beige),
                                ('FONTNAME', (0,1), (-1,-1), chinese_font_name if chinese_available else 'Helvetica'),
                                ('FONTSIZE', (0,1), (-1,-1), 10),
                                ('GRID',(0,0),(-1,-1),1,colors.black)
                            ]))
                            elements.append(table)
                            elements.append(Spacer(1, 12))
                            
                    elif element.name == 'pre':
                        # 代码块
                        text = element.get_text().strip()
                        if text:
                            # 将代码分行处理
                            for line in text.split('\n'):
                                elements.append(Paragraph(line or " ", code_style))
                            elements.append(Spacer(1, 6))
                            
                    elif element.name == 'hr':
                        # 分隔线
                        elements.append(Spacer(1, 12))
                        elements.append(Paragraph("─" * 50, normal_style))
                        elements.append(Spacer(1, 12))
                        
                    else:
                        # 其他元素当作段落处理
                        text = element.get_text().strip()
                        if text:
                            elements.append(Paragraph(text, normal_style))
                            elements.append(Spacer(1, 6))
                            
        except Exception as e:
            # 如果HTML解析失败，添加错误信息
            elements.append(Paragraph(f"HTML解析失败: {str(e)}", base_style))
            
        return elements
    
    @staticmethod
    def _process_inline_html_elements(element):
        """处理段落中的内联HTML元素（粗体、斜体等）"""
        text = ""
        for content in element.contents:
            if hasattr(content, 'name'):
                if content.name == 'strong' or content.name == 'b':
                    text += f"<b>{content.get_text()}</b>"
                elif content.name == 'em' or content.name == 'i':
                    text += f"<i>{content.get_text()}</i>"
                elif content.name == 'code':
                    text += f"<font name='Courier'>{content.get_text()}</font>"
                else:
                    text += content.get_text()
            else:
                text += str(content)
        return text.strip()
    
    @staticmethod
    def export_task_result_to_markdown(task_result, output_path=None, format_type='preview'):
        """导出任务结果为Markdown格式
        
        Args:
            task_result: TaskResult对象
            output_path: 输出文件路径
            format_type: 导出格式类型 ('raw' | 'preview')
                - 'raw': 纯Markdown源码格式
                - 'preview': HTML预览格式（渲染后的Markdown）
        """
        try:
            if not output_path:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                file_ext = '.md' if format_type == 'raw' else '.html'
                output_path = f"task_result_{task_result.task_id}_{timestamp}{file_ext}"
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # 获取任务结果内容
            content = task_result.answer or "暂无处理结果"
            content = str(content).strip()
            
            if format_type == 'raw':
                # 原始Markdown格式输出
                # 清理内容，确保是有效的Markdown格式
                cleaned_content = DocumentService._clean_markdown_content(content)
                
                # 直接写入.md文件
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(cleaned_content)
                
                try:
                    current_app.logger.info(f"Markdown原始格式导出成功: {output_path}")
                except RuntimeError:
                    print(f"Markdown原始格式导出成功: {output_path}")
                
                return True, output_path, None
                
            else:  # format_type == 'preview'
                # HTML预览格式输出
                # 清理内容后再转换
                cleaned_content = DocumentService._clean_markdown_content(content)
                html_content = DocumentService._convert_markdown_to_html(cleaned_content)
                
                # 生成完整的HTML页面
                full_html = DocumentService._create_html_page(html_content, f"任务结果 - {task_result.task_id}")
                
                # 写入HTML文件
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(full_html)
                
                try:
                    current_app.logger.info(f"Markdown预览格式导出成功: {output_path}")
                except RuntimeError:
                    print(f"Markdown预览格式导出成功: {output_path}")
                
                return True, output_path, None
                
        except Exception as e:
            error_msg = f"Markdown导出失败: {str(e)}"
            try:
                current_app.logger.error(error_msg, exc_info=True)
            except RuntimeError:
                print(error_msg)
                import traceback
                traceback.print_exc()
            return False, None, error_msg
    
    @staticmethod
    def _convert_markdown_to_html(markdown_content):
        """将Markdown内容转换为HTML"""
        try:
            import markdown
            
            # 配置Markdown扩展
            extensions = [
                'markdown.extensions.extra',  # 包含tables, fenced_code等
                'markdown.extensions.codehilite',  # 代码高亮
                'markdown.extensions.toc',  # 目录
                'markdown.extensions.nl2br',  # 换行转换
            ]
            
            # 转换Markdown为HTML
            md = markdown.Markdown(extensions=extensions)
            html_content = md.convert(markdown_content)
            
            return html_content
            
        except ImportError:
            # 如果markdown库不可用，返回基本的HTML格式
            escaped_content = (markdown_content
                             .replace('&', '&amp;')
                             .replace('<', '&lt;')
                             .replace('>', '&gt;')
                             .replace('\n', '<br>'))
            return f"<div><pre>{escaped_content}</pre></div>"
        except Exception as e:
            # 如果Markdown转换失败，返回纯文本的HTML格式
            escaped_content = (markdown_content
                             .replace('&', '&amp;')
                             .replace('<', '&lt;')
                             .replace('>', '&gt;')
                             .replace('\n', '<br>'))
            return f"<div><pre>{escaped_content}</pre></div>"
    
    @staticmethod
    def _create_pdf_html_page(html_content, title="任务结果"):
        """创建适合PDF导出的HTML页面（优化打印样式）"""
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        @page {{
            margin: 2cm;
            size: A4;
        }}
        
        body {{
            font-family: 'SimSun', '宋体', Arial, sans-serif;
            font-size: 12pt;
            line-height: 1.6;
            color: #333;
            margin: 0;
            padding: 0;
        }}
        
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 1.5em;
            margin-bottom: 0.8em;
            font-weight: bold;
            break-after: avoid;
        }}
        
        h1 {{
            font-size: 18pt;
            text-align: center;
            border-bottom: 2px solid #333;
            padding-bottom: 0.5em;
        }}
        
        h2 {{ font-size: 16pt; }}
        h3 {{ font-size: 14pt; }}
        h4 {{ font-size: 13pt; }}
        h5, h6 {{ font-size: 12pt; }}
        
        p {{
            margin-bottom: 1em;
            text-align: justify;
        }}
        
        pre, code {{
            font-family: 'Courier New', monospace;
            background-color: #f5f5f5;
            border: 1px solid #ddd;
        }}
        
        pre {{
            padding: 10pt;
            margin: 1em 0;
            overflow: hidden;
            border-radius: 4px;
            break-inside: avoid;
        }}
        
        code {{
            padding: 2pt 4pt;
            border-radius: 2px;
        }}
        
        blockquote {{
            border-left: 4px solid #ddd;
            margin: 1em 0;
            padding-left: 1em;
            color: #666;
            font-style: italic;
        }}
        
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
            break-inside: avoid;
        }}
        
        th, td {{
            border: 1px solid #ddd;
            padding: 8pt;
            text-align: left;
        }}
        
        th {{
            background-color: #f9f9f9;
            font-weight: bold;
        }}
        
        ul, ol {{
            margin: 1em 0;
            padding-left: 2em;
        }}
        
        li {{
            margin-bottom: 0.5em;
        }}
        
        .header {{
            text-align: center;
            margin-bottom: 2em;
            padding-bottom: 1em;
            border-bottom: 1px solid #ccc;
        }}
        
        .export-info {{
            font-size: 10pt;
            color: #666;
            margin-top: 0.5em;
        }}
        
        .footer {{
            margin-top: 2em;
            padding-top: 1em;
            border-top: 1px solid #ccc;
            text-align: center;
            font-size: 10pt;
            color: #666;
        }}
        
        /* 确保内容不会被截断 */
        * {{
            box-sizing: border-box;
        }}
        
        img {{
            max-width: 100%;
            height: auto;
        }}
        
        /* 避免分页时内容断裂 */
        h1, h2, h3, h4, h5, h6, p, li, td, th {{
            break-inside: avoid;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{title}</h1>
        <div class="export-info">
            导出时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}
        </div>
    </div>
    
    <div class="content">
        {html_content}
    </div>
    
    <div class="footer">
        <p>此文档由系统自动生成 · Markdown格式预览</p>
    </div>
 </body>
 </html>"""
    
    @staticmethod 
    def _create_html_page(html_content, title="任务结果"):
        """创建完整的HTML页面（用于网页预览）"""
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background-color: #fff;
        }}
        
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: 600;
            line-height: 1.25;
        }}
        
        h1 {{
            padding-bottom: 0.3em;
            border-bottom: 1px solid #eaecef;
        }}
        
        p {{
            margin-bottom: 16px;
        }}
        
        code {{
            background-color: #f1f3f4;
            border-radius: 3px;
            font-size: 85%;
            margin: 0;
            padding: 0.2em 0.4em;
        }}
        
        pre {{
            background-color: #f6f8fa;
            border-radius: 6px;
            font-size: 85%;
            line-height: 1.45;
            overflow: auto;
            padding: 16px;
        }}
        
        pre code {{
            background-color: transparent;
            border: 0;
            display: inline;
            line-height: inherit;
            margin: 0;
            max-width: auto;
            overflow: visible;
            padding: 0;
            word-wrap: normal;
        }}
        
        blockquote {{
            border-left: 4px solid #dfe2e5;
            margin: 0;
            padding: 0 16px;
            color: #6a737d;
        }}
        
        table {{
            border-collapse: collapse;
            border-spacing: 0;
            width: 100%;
            margin-bottom: 16px;
        }}
        
        table th, table td {{
            border: 1px solid #dfe2e5;
            padding: 6px 13px;
        }}
        
        table th {{
            background-color: #f6f8fa;
            font-weight: 600;
        }}
        
        ul, ol {{
            padding-left: 2em;
            margin-bottom: 16px;
        }}
        
        li {{
            margin-bottom: 0.25em;
        }}
        
        .header {{
            text-align: center;
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 2px solid #e1e4e8;
        }}
        
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #e1e4e8;
            text-align: center;
            color: #586069;
            font-size: 14px;
        }}
        
        @media print {{
            body {{
                max-width: none;
                margin: 0;
                padding: 15px;
            }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{title}</h1>
        <p>导出时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}</p>
    </div>
    
    <div class="content">
        {html_content}
    </div>
    
    <div class="footer">
        <p>此文档由系统自动生成</p>
    </div>
</body>
</html>"""
    
    @staticmethod
    def _export_text_to_pdf_legacy(task_result, output_path):
        """原有方法：使用ReportLab直接处理文本（作为备用方案）"""
        try:
            # 注册中文字体（使用改进的策略）
            chinese_available = False
            chinese_font_name = 'ChineseFont'
            
            try:
                # 尝试多种字体注册策略
                font_registered = False
                
                # 策略1: 优先尝试系统字体（更可靠）
                font_path = DocumentService._find_chinese_font()
                if font_path:
                    try:
                        current_app.logger.info(f"正在注册系统字体: {font_path}")
                    except RuntimeError:
                        print(f"正在注册系统字体: {font_path}")
                    
                    try:
                        # 尝试标准TTF注册
                        pdfmetrics.registerFont(TTFont(chinese_font_name, font_path))
                        chinese_available = True
                        font_registered = True
                    except Exception as ttf_error:
                        try:
                            # 如果是TTC文件，尝试使用子字体
                            if font_path.endswith('.ttc'):
                                pdfmetrics.registerFont(TTFont(chinese_font_name, font_path, subfontIndex=0))
                                chinese_available = True
                                font_registered = True
                            else:
                                raise ttf_error
                        except:
                            # 最后尝试强制使用subfontIndex
                            try:
                                pdfmetrics.registerFont(TTFont(chinese_font_name, font_path, subfontIndex=0))
                                chinese_available = True
                                font_registered = True
                            except:
                                pass
                
                # 策略2: 如果系统字体失败，尝试内置CID字体
                if not font_registered:
                    try:
                        # 使用ReportLab的内置字体作为备选
                        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                        pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
                        chinese_font_name = 'STSong-Light'
                        chinese_available = True
                        font_registered = True
                        try:
                            current_app.logger.info("使用内置中文字体: STSong-Light")
                        except RuntimeError:
                            print("使用内置中文字体: STSong-Light")
                    except:
                        # 尝试其他内置CID字体
                        cid_fonts = ['STSongStd-Light', 'HeiseiMin-W3', 'HeiseiKakuGo-W5', 'MSung-Light', 'MSungStd-Light']
                        for cid_font in cid_fonts:
                            try:
                                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                                pdfmetrics.registerFont(UnicodeCIDFont(cid_font))
                                chinese_font_name = cid_font
                                chinese_available = True
                                font_registered = True
                                try:
                                    current_app.logger.info(f"使用内置中文字体: {cid_font}")
                                except RuntimeError:
                                    print(f"使用内置中文字体: {cid_font}")
                                break
                            except:
                                continue
                
                # 策略3: 如果都失败，使用Helvetica但标记为支持中文（降级方案）
                if not font_registered:
                    try:
                        current_app.logger.warning("使用Helvetica作为备选字体")
                    except RuntimeError:
                        print("使用Helvetica作为备选字体")
                    chinese_font_name = 'Helvetica'
                    chinese_available = False  # 不支持中文，但不会报错
                
                if font_registered:
                    try:
                        current_app.logger.info(f"中文字体注册成功: {chinese_font_name}")
                    except RuntimeError:
                        print(f"中文字体注册成功: {chinese_font_name}")
                        
            except Exception as e:
                try:
                    current_app.logger.error(f"字体注册过程失败: {str(e)}")
                except RuntimeError:
                    print(f"字体注册过程失败: {str(e)}")
                chinese_font_name = 'Helvetica'
                chinese_available = False
            
            # 创建PDF文档
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # 创建自定义样式
            if chinese_available:
                title_style = ParagraphStyle(
                    'ChineseTitle',
                    parent=styles['Title'],
                    fontName=chinese_font_name,
                    fontSize=16,
                    spaceAfter=20
                )
                normal_style = ParagraphStyle(
                    'ChineseNormal',
                    parent=styles['Normal'],
                    fontName=chinese_font_name,
                    fontSize=12,
                    spaceAfter=12
                )
            else:
                title_style = styles['Title']
                normal_style = styles['Normal']
            
            # 构建文档内容
            story = []
            
            # 只导出处理结果，不包含任务信息和元数据
            if task_result.answer:
                # 首先使用Markdown内容清理（移除代码块标记等）
                cleaned_markdown = DocumentService._clean_markdown_content(task_result.answer)
                
                # 将Markdown转换为HTML预览格式
                try:
                    html_content = DocumentService._convert_markdown_to_html(cleaned_markdown)
                    
                    # 将HTML转换为ReportLab可用的格式
                    story.extend(DocumentService._convert_html_to_reportlab_elements(
                        html_content, normal_style, chinese_available, chinese_font_name
                    ))
                    
                    try:
                        current_app.logger.info("✅ 使用Markdown预览格式转换PDF")
                    except RuntimeError:
                        print("✅ 使用Markdown预览格式转换PDF")
                        
                except Exception as html_error:
                    # 如果HTML转换失败，回退到纯文本方式
                    try:
                        current_app.logger.warning(f"HTML转换失败，使用纯文本模式: {html_error}")
                    except RuntimeError:
                        print(f"HTML转换失败，使用纯文本模式: {html_error}")
                    
                    answer_text = DocumentService._clean_text_for_pdf(cleaned_markdown)
                    
                    # 将文本按行分割并添加到PDF
                    answer_lines = answer_text.split('\n')
                    for line in answer_lines:
                        if line.strip():
                            # 对特殊字符进行HTML转义处理
                            escaped_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            story.append(Paragraph(escaped_line, normal_style))
                        else:
                            # 空行添加间距
                            story.append(Spacer(1, 6))
            else:
                # 如果没有结果内容，添加一个提示
                story.append(Paragraph("暂无处理结果", normal_style))
            
            # 生成PDF
            doc.build(story)
            
            try:
                current_app.logger.info(f"PDF导出成功（使用原有方法）: {output_path}")
            except RuntimeError:
                print(f"PDF导出成功（使用原有方法）: {output_path}")
            return True, output_path, None
            
        except Exception as e:
            error_msg = f"PDF导出失败（原有方法）: {str(e)}"
            try:
                current_app.logger.error(error_msg, exc_info=True)
            except RuntimeError:
                print(error_msg)
                import traceback
                traceback.print_exc()
            return False, None, error_msg
    
    @staticmethod
    def _clean_text_for_pdf(text):
        """清理和预处理文本以确保PDF正确显示"""
        if not text:
            return ""
        
        # 确保是字符串
        text = str(text)
        
        # 替换可能导致问题的字符
        replacements = {
            '\u2018': "'",  # 左单引号
            '\u2019': "'",  # 右单引号
            '\u201c': '"',  # 左双引号
            '\u201d': '"',  # 右双引号
            '\u2013': '-',  # en dash
            '\u2014': '--', # em dash
            '\u2026': '...', # 省略号
            '\u00a0': ' ',  # 不间断空格
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # 确保换行符统一
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # 移除可能导致问题的控制字符
        import re
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        
        return text

    @staticmethod
    def _find_chinese_font():
        """查找中文字体"""
        font_paths = [
            # 我们脚本安装的字体目录（最优先）
            '/usr/share/fonts/truetype/chinese/wqy-microhei.ttc',
            '/usr/share/fonts/truetype/chinese/NotoSansCJK.ttc',
            '/usr/share/fonts/truetype/chinese/NotoSansCJK.otf',
            '/usr/share/fonts/truetype/chinese/wqy-zenhei.ttc',
            # Windows - 优先使用.ttf文件，避免.ttc文件的兼容问题
            'C:/Windows/Fonts/simhei.ttf',  # 黑体
            'C:/Windows/Fonts/simsun.ttf',  # 宋体（如果有单独的ttf文件）
            'C:/Windows/Fonts/simkai.ttf',  # 楷体
            'C:/Windows/Fonts/msyh.ttf',    # 微软雅黑（如果有单独的ttf文件）
            'C:/Windows/Fonts/arial.ttf',   # Arial作为备选
            # 如果没有单独的ttf文件，再尝试ttc文件
            'C:/Windows/Fonts/simsun.ttc',
            'C:/Windows/Fonts/msyh.ttc',
            # macOS  
            '/System/Library/Fonts/Helvetica.ttc',
            '/Library/Fonts/Arial Unicode.ttf',
            '/System/Library/Fonts/PingFang.ttc',
            # Linux - 大幅扩展Linux字体路径
            '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
            '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc',
            '/usr/share/fonts/truetype/arphic/uming.ttc',
            '/usr/share/fonts/truetype/arphic/ukai.ttc',
            '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            # 其他可能的Linux路径
            '/usr/share/fonts/TTF/wqy-microhei.ttc',
            '/usr/share/fonts/wqy-microhei/wqy-microhei.ttc',
            '/usr/share/fonts/wenquanyi/wqy-microhei/wqy-microhei.ttc'
        ]
        
        for path in font_paths:
            if os.path.exists(path):
                try:
                    current_app.logger.info(f"找到字体文件: {path}")
                except RuntimeError:
                    # 在没有应用上下文时使用print
                    print(f"找到字体文件: {path}")
                return path
        
        # 如果静态路径都找不到，尝试使用系统命令查找
        try:
            import subprocess
            result = subprocess.run(['fc-list', ':lang=zh', 'file'], 
                                  capture_output=True, text=True, timeout=10)
            if result.returncode == 0 and result.stdout:
                # 解析fc-list输出，获取第一个字体文件路径
                lines = result.stdout.strip().split('\n')
                for line in lines:
                    if ':' in line:
                        font_path = line.split(':')[0].strip()
                        if os.path.exists(font_path):
                            try:
                                current_app.logger.info(f"通过fc-list找到字体: {font_path}")
                            except RuntimeError:
                                print(f"通过fc-list找到字体: {font_path}")
                            return font_path
        except Exception as e:
            try:
                current_app.logger.warning(f"fc-list命令执行失败: {e}")
            except RuntimeError:
                print(f"fc-list命令执行失败: {e}")
        
        try:
            current_app.logger.warning("未找到合适的中文字体文件")
        except RuntimeError:
            print("未找到合适的中文字体文件")
        return None
    
    @staticmethod
    def get_file_stats(file_path):
        """获取文件统计信息"""
        try:
            if not os.path.exists(file_path):
                return None
            
            stat = os.stat(file_path)
            file_extension = os.path.splitext(file_path)[1].lower().lstrip('.')
            
            return {
                'size': stat.st_size,
                'size_human': DocumentService._human_readable_size(stat.st_size),
                'modified_time': datetime.fromtimestamp(stat.st_mtime),
                'created_time': datetime.fromtimestamp(stat.st_ctime),
                'extension': file_extension,
                'can_preview': file_extension in DocumentService.PREVIEW_SUPPORTED_TYPES
            }
            
        except Exception as e:
            current_app.logger.error(f"获取文件统计失败: {file_path} - 错误: {str(e)}")
            return None
    
    @staticmethod
    def _human_readable_size(size_bytes):
        """将字节数转换为人类可读的格式"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB" 

    @staticmethod
    def export_task_results_to_excel(items_data, task_info, output_path=None):
        """导出任务分页结果为Excel文件"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            from datetime import datetime
            
            if not output_path:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = f"task_results_{task_info.get('id', 'unknown')}_{timestamp}.xlsx"
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # 创建工作簿
            wb = Workbook()
            ws = wb.active
            ws.title = "任务结果"
            
            # 定义样式
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            
            content_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            border = Border(
                left=Side(style='thin'), 
                right=Side(style='thin'), 
                top=Side(style='thin'), 
                bottom=Side(style='thin')
            )
            
            # 写入标题信息
            ws.merge_cells('A1:E1')
            title_cell = ws['A1']
            title_cell.value = f"任务结果导出 - {task_info.get('title', '未命名任务')}"
            title_cell.font = Font(size=16, bold=True)
            title_cell.alignment = Alignment(horizontal="center")
            
            # 写入任务基本信息
            info_row = 3
            ws[f'A{info_row}'] = "任务类型："
            ws[f'B{info_row}'] = task_info.get('task_type_display', task_info.get('task_type', ''))
            ws[f'D{info_row}'] = "创建时间："
            ws[f'E{info_row}'] = task_info.get('created_at', '')[:19] if task_info.get('created_at') else ''
            
            info_row += 1
            ws[f'A{info_row}'] = "任务状态："
            ws[f'B{info_row}'] = task_info.get('status_display', task_info.get('status', ''))
            ws[f'D{info_row}'] = "导出时间："
            ws[f'E{info_row}'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            # 根据任务类型定义不同的列配置
            task_type = task_info.get('task_type', 'standard_review')
            
            if task_type == 'standard_recommendation':
                # 标准推荐
                headers = ['排序序号', '项目名称', '原文内容', '参考标准']
                field_keys = ['sn', 'projectName', 'originalText', 'referenceStandard']
                column_widths = [10, 25, 50, 40]
                long_text_columns = [3, 4]  # 原文内容、参考标准
            elif task_type == 'standard_compliance':
                # 标准符合性检查
                headers = ['排序序号', '项目名称', '原文内容', '是否符合标准', '建议改写内容', '参考标准']
                field_keys = ['sn', 'projectName', 'originalText', 'isCompliant', 'suggestedRewrite', 'referenceStandard']
                column_widths = [10, 25, 40, 15, 40, 35]
                long_text_columns = [3, 5, 6]  # 原文内容、建议改写内容、参考标准
            else:
                # 标准审查（默认）
                headers = ['序号', '问题位置', '原文', '问题描述', '修改建议']
                field_keys = ['sn', 'issueLocation', 'originalText', 'issueDescription', 'recommendedModification']
                column_widths = [8, 20, 40, 30, 40]
                long_text_columns = [3, 4, 5]  # 原文、问题描述、修改建议
            
            # 写入表头（从第6行开始）
            header_row = 6
            
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=header_row, column=col)
                cell.value = header
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = border
            
            # 设置列宽
            for col, width in enumerate(column_widths, 1):
                ws.column_dimensions[get_column_letter(col)].width = width
            
            # 写入数据
            data_start_row = header_row + 1
            for row_idx, item in enumerate(items_data, 1):
                current_row = data_start_row + row_idx - 1
                
                # 根据字段配置写入数据到各列
                data_values = []
                for field_key in field_keys:
                    value = item.get(field_key, '')
                    # 对于序号字段，如果为空则使用行号
                    if field_key == 'sn' and not value:
                        value = row_idx
                    data_values.append(value)
                
                for col, value in enumerate(data_values, 1):
                    cell = ws.cell(row=current_row, column=col)
                    cell.value = str(value) if value is not None else ''
                    cell.alignment = content_alignment
                    cell.border = border
                    
                    # 为内容较长的列设置行高
                    if col in long_text_columns:
                        # 根据文本长度动态调整行高
                        text_length = len(str(value)) if value else 0
                        if text_length > 50:
                            # 估算需要的行高（每50个字符大约需要15磅高度）
                            estimated_height = min(max(15, (text_length // 50) * 15), 100)
                            if ws.row_dimensions[current_row].height is None or ws.row_dimensions[current_row].height < estimated_height:
                                ws.row_dimensions[current_row].height = estimated_height
            
            # 冻结窗格（冻结表头）
            ws.freeze_panes = f'A{data_start_row}'
            
            # 添加筛选功能
            if items_data:
                last_column = get_column_letter(len(headers))
                ws.auto_filter.ref = f'A{header_row}:{last_column}{data_start_row + len(items_data) - 1}'
            
            # 保存文件
            wb.save(output_path)
            
            try:
                current_app.logger.info(f"Excel导出成功: {output_path}, 共导出 {len(items_data)} 条记录")
            except RuntimeError:
                print(f"Excel导出成功: {output_path}, 共导出 {len(items_data)} 条记录")
            
            return True, output_path, None
            
        except Exception as e:
            error_msg = f"Excel导出失败: {str(e)}"
            try:
                current_app.logger.error(error_msg, exc_info=True)
            except RuntimeError:
                print(error_msg)
                import traceback
                traceback.print_exc()
            return False, None, error_msg